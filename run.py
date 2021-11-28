from docx import Document
from docx.oxml.ns import qn

import os
import pathlib
from datetime import datetime, date

dir_path = pathlib.Path(__file__).parent
docx_path = dir_path / "all.docx"
template_path = dir_path / "template.docx"
txt_path = dir_path / "all.txt"
txt_sep_list = dir_path/"分割"
txt_sep_list = list(txt_sep_list.glob("*"))
csv_path = dir_path / "statistic.csv"

yyyymmdd = date.today().strftime("%Y/%m/%d")

def convert_to_txt():
  document = Document(docx_path)
  with open(txt_path, "w", encoding='utf-8') as f:
    for para in document.paragraphs:
      print(para.text, file=f)

def convert_to_all_txt():
  with open(txt_path, "w", encoding='utf-8') as all:
    for i in txt_sep_list:
      with open(i,"r",encoding='utf-8') as sep:
        data = sep.read()
        print(data, file=all)

def convert_to_doc():
  document = Document(template_path)
  with open(txt_path, "r", encoding='utf-8') as f:
    txt_content = f.read().split("\n")
    for para_content in txt_content:
      para = document.add_paragraph()
      run = para.add_run(para_content)
      run.font.name = "源暎こぶり明朝 v6"
      run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
    document.save(docx_path)

def get_num_of_char():
  num_of_char = 0
  document = Document(docx_path)
  for para in document.paragraphs:
    num_of_char += len(para.text)
    "{:,}".format(num_of_char)
  with open(csv_path, "a") as f:
    add_content = f'{yyyymmdd},{num_of_char}'
    print(add_content, file=f)

txt_mtime = os.path.getmtime(txt_path)
docx_mtime = os.path.getmtime(docx_path)

'''
if txt_mtime >= docx_mtime:
  #convert_to_doc()
else:
  #convert_to_txt()
'''

#convert_to_txt()
#convert_to_all_txt()
#convert_to_doc()

get_num_of_char()